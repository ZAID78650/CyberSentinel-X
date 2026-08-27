"""Generate the CyberSentinel-X SIH 2026 project document (.docx).

Run from the repository root with the backend virtualenv python:

    cd backend && .venv/bin/python ../scripts/generate_sih_docx.py

Output: docs/CyberSentinel-X_SIH_2026.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "CyberSentinel-X_SIH_2026.docx"
SHOTS = ROOT / "docs" / "screenshots"

NAVY = RGBColor(0x0B, 0x3D, 0x91)      # headings
STEEL = RGBColor(0x1F, 0x38, 0x64)     # title page
ACCENT = RGBColor(0x8B, 0x1A, 0x1A)    # small accent
GRAY = RGBColor(0x59, 0x59, 0x59)


# --------------------------------------------------------------------------
# Low-level helpers
# --------------------------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Right-click here and choose \u201cUpdate Field\u201d to generate the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(t)
    run._r.append(fld_end)
    run.font.size = Pt(10)
    run.font.italic = True


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = GRAY


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if isinstance(it, tuple):
            bold, rest = it
            r = p.add_run(bold)
            r.bold = True
            p.add_run(rest)
        else:
            p.add_run(it)


def make_table(doc, headers, rows, widths=None, header_fill="0B3D91"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        set_cell_bg(hdr[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
            if i == 0:
                r.bold = True
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    return p


# --------------------------------------------------------------------------
# Document setup
# --------------------------------------------------------------------------

doc = Document()

# Base styles
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = NAVY
    st.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
    st.paragraph_format.space_after = Pt(6)

sec = doc.sections[0]
sec.page_width = Inches(8.27)   # A4
sec.page_height = Inches(11.69)
for attr in ("top_margin", "bottom_margin"):
    setattr(sec, attr, Inches(0.9))
for attr in ("left_margin", "right_margin"):
    setattr(sec, attr, Inches(1.0))

add_page_number_footer(sec)

props = doc.core_properties
props.title = "CyberSentinel-X — Smart India Hackathon 2026"
props.subject = "Agentic AI-Powered Autonomous Cyber Threat Detection, Investigation & Response Platform"
props.author = "Team CyberSentinel-X"


# --------------------------------------------------------------------------
# 1. Cover page
# --------------------------------------------------------------------------

for _ in range(5):
    spacer(doc, 12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CYBERSENTINEL-X")
r.font.size = Pt(40)
r.font.bold = True
r.font.color.rgb = STEEL

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Agentic AI-Powered Autonomous Cyber Threat Detection,\nInvestigation & Response Platform")
r.font.size = Pt(15)
r.font.color.rgb = NAVY

spacer(doc, 10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Smart India Hackathon 2026 — Project Submission")
r.font.size = Pt(13)
r.font.bold = True
r.font.color.rgb = ACCENT

spacer(doc, 6)

cover_shot = SHOTS / "02-dashboard.png"
if cover_shot.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(cover_shot), width=Inches(6.1))

spacer(doc, 10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Team CyberSentinel-X  ·  Smart India Hackathon 2026  ·  Ministry of Electronics & IT (MeitY)")
r.font.size = Pt(11)
r.font.color.rgb = GRAY

doc.add_page_break()

# --------------------------------------------------------------------------
# 2. Table of contents
# --------------------------------------------------------------------------

doc.add_heading("Table of Contents", level=1)
add_toc(doc)
doc.add_page_break()

# --------------------------------------------------------------------------
# 3. The Idea
# --------------------------------------------------------------------------

doc.add_heading("1. The Idea of CyberSentinel-X", level=1)

doc.add_paragraph(
    "CyberSentinel-X is a next-generation, AI-powered Security Operations Center (SOC) platform that "
    "turns raw security telemetry into a fully autonomous defense pipeline: detect, investigate, "
    "correlate, map to MITRE ATT&CK, score risk, recommend responses, and document everything — "
    "all driven by a controlled team of specialized AI agents, with a human always in the loop for "
    "high-impact decisions."
)

doc.add_paragraph(
    "Modern cyberattacks are too fast, too automated, and too numerous for manual triage. Security teams "
    "suffer from alert fatigue, long mean-time-to-detect (MTTD) and mean-time-to-respond (MTTR), and a "
    "chronic shortage of skilled analysts. CyberSentinel-X attacks this problem at its root: instead of "
    "showing analysts a wall of raw alerts, it autonomously investigates each incident, produces "
    "evidence-backed explanations, reconstructs the attack path, predicts what the adversary does next, "
    "and recommends the cheapest effective response — all in real time and fully auditable."
)

p = doc.add_paragraph()
r = p.add_run("What makes the idea distinctive: ")
r.bold = True
p.add_run("it is an agentic system built for production trust. Every AI decision is explainable, every "
          "agent action is restricted to allowlisted read-only tools, no response executes without human "
          "approval, and forensic evidence is preserved on a tamper-evident, blockchain-style ledger. The "
          "platform therefore demonstrates not just an AI that can detect attacks, but an AI that can be "
          "trusted and safely deployed inside a real organization.")

doc.add_heading("Why this matters", level=2)
bullets(doc, [
    ("Hybrid detection — ", "deterministic rules, Isolation-Forest anomaly detection, threat intelligence "
     "and AI reasoning work together; AI is never the only decision-maker, so it never hallucinates a false positive."),
    ("Autonomous triage — ", "the AI pipeline does in seconds what a junior analyst takes hours to do: "
     "correlate events, gather evidence, assign a verdict and a confidence score."),
    ("Explainable, not black-box — ", "risk scores come with a weighted factor breakdown and every "
     "investigation cites the exact evidence behind its verdict."),
    ("Safe by construction — ", "agents only call allowlisted database tools, responses are simulated, and "
     "every high-impact action needs a human approval."),
])

# --------------------------------------------------------------------------
# 4. Implementation for SIH
# --------------------------------------------------------------------------

doc.add_heading("2. Implementation for Smart India Hackathon 2026", level=1)

doc.add_paragraph(
    "CyberSentinel-X was built from scratch as a fully working, end-to-end prototype targeting the "
    "cybersecurity problem statement of SIH 2026. It is not a mockup: the frontend, backend, machine "
    "learning models, agent orchestration, database, and reporting are all real and runnable. The full "
    "stack ships with Docker Compose, an automated CI/CD pipeline, 34 automated backend tests, and "
    "database migrations — a production-ready posture rather than a demo-only scaffold."
)

doc.add_heading("Problem statement addressed", level=2)
doc.add_paragraph(
    "Building an autonomous, AI-driven threat detection and response capability for a Security Operations "
    "Center — reducing alert fatigue, improving detection speed and accuracy, and providing explainable, "
    "actionable intelligence to human analysts, all with robust security, privacy, and auditability."
)

doc.add_heading("Technology stack", level=2)
make_table(doc,
    ["Layer", "Technology"],
    [
        ["Frontend", "React + Vite + TypeScript + Tailwind CSS, React Flow (attack graphs), Recharts (analytics), WebSockets for real-time streaming"],
        ["Backend", "FastAPI (Python 3.12), async REST + WebSocket, Pydantic validation, SQLAlchemy 2.0"],
        ["Database", "PostgreSQL (Docker) / SQLite (local dev), Alembic migrations, 25+ tables"],
        ["Machine Learning", "Isolation Forest anomaly detection, population-stability model-drift monitoring, UEBA behavioral baselines"],
        ["AI / Agents", "Agent orchestration state machine (Detection → Investigation → Threat Intel → Risk → Response), pluggable LLM providers (local / OpenAI / Gemini)"],
        ["Knowledge", "RAG pipeline with a local vector store (numpy backend, Chroma optional) over playbooks, policies, CVEs and MITRE references"],
        ["Threat Intelligence", "Local STIX-shaped feed + adapter interface (swappable for live TAXII / vendor APIs), MITRE ATT&CK dataset"],
        ["Forensics", "SHA-256 evidence hashing → Merkle tree roots → proof-of-work style blocks on a permissioned local ledger"],
        ["Ops", "Docker Compose, GitHub Actions CI, render.yaml deployment, pytest suite, smoke tests"],
    ],
    widths=[1.3, 4.9],
)

doc.add_heading("What was actually built", level=2)
bullets(doc, [
    "A real-time SOC console with Dashboard, Live Events, Alerts, Incidents, Campaigns, Attack Graph, Risk Overview, Response Center, Human Approvals, Actions Log, Incident Reports and more (40 pages).",
    "A controlled agentic AI pipeline that autonomously investigates incidents, maps techniques to MITRE ATT&CK, builds attack graphs, and writes automated HTML + PDF incident reports.",
    "An explainable risk engine scoring incidents out of 100 using five weighted, evidence-backed factors.",
    "A layered web-application defense center with request-ID tracking, rate limiting, payload inspection, malware hash blocking and brute-force protection.",
    "RBAC with three roles (ADMIN, SECURITY_ANALYST, VIEWER), JWT access/refresh tokens, bcrypt hashing, and a full audit trail of every action.",
    "A synthetic attack simulator (Account Takeover, Brute Force, Malware, Data Exfiltration, Privilege Escalation) that generates correlated event streams for live demonstrations.",
    "A tamper-evident evidence ledger with Merkle-tree verification and a one-click tamper test for judges.",
])

doc.add_heading("Live demo flow for judges", level=2)
doc.add_paragraph(
    "The 5-minute judge demonstration walks through: login (with optional SSO) → Judge Mode overview of the "
    "whole pipeline with real counts → a live simulated attack streaming over WebSocket → AI investigation "
    "with verdict and confidence → threat intelligence and MITRE mapping → attack graph → explainable risk "
    "score → human-approved simulated response → automated incident report → blockchain evidence "
    "verification. Every number shown is computed from real platform state; simulated data, dataset data "
    "(UNSW-NB15 — see Section 8), and model predictions are honestly labelled with provenance badges."
)

# --------------------------------------------------------------------------
# 5. Introduction
# --------------------------------------------------------------------------

doc.add_heading("3. Introduction", level=1)

doc.add_paragraph(
    "Cybersecurity is one of the defining challenges of the digital age. Governments, banks, hospitals, "
    "power grids and enterprises are under continuous assault from ransomware gangs, nation-state actors, "
    "and financially motivated criminals. The attack surface keeps growing — cloud workloads, IoT devices, "
    "remote work, and software supply chains — while the number of skilled analysts cannot keep pace. "
    "Traditional security tools generate overwhelming volumes of alerts, most of which are false positives, "
    "and the true incidents that matter are buried under the noise."
)

doc.add_paragraph(
    "CyberSentinel-X is our answer: a Security Operations Center that thinks like a senior analyst. It "
    "ingests security events from any source, detects suspicious behavior with a hybrid of deterministic "
    "rules, machine learning and threat intelligence, and then deploys a team of specialized AI agents to "
    "investigate what it found. The investigation is not a black box — it produces a timeline, evidence "
    "list, MITRE technique mappings, an attack graph, and a confidence-scored verdict. A risk engine "
    "quantifies the business impact, and a response agent recommends concrete containment actions that a "
    "human analyst reviews and approves. Finally, every step is recorded in an auditable, tamper-evident "
    "ledger and summarized in an automated incident report."
)

doc.add_paragraph(
    "The result is a platform that amplifies human analysts rather than replacing them: it handles the "
    "tedious investigation work at machine speed, presents findings that are explainable and verifiable, "
    "and keeps the final decision with a human. This human-in-the-loop autonomy is what makes "
    "CyberSentinel-X practical for real, high-stakes deployments."
)

# --------------------------------------------------------------------------
# 6. Core Fundamentals
# --------------------------------------------------------------------------

doc.add_heading("4. Core Fundamentals", level=1)

doc.add_paragraph(
    "The platform is built on eight core fundamentals. Together they define how CyberSentinel-X thinks, "
    "why its outputs are trustworthy, and how it stays safe to deploy."
)

core = [
    ("4.1  Hybrid Detection Engine",
     "Detection never relies on a single signal. A deterministic rules engine captures known-bad patterns, "
     "an Isolation Forest model flags statistical anomalies in event streams, a threat-intelligence layer "
     "matches indicators of compromise, and an AI agent correlates the results. This hybrid design keeps "
     "precision high (fewer false positives) and recall high (novel attacks are still caught by the anomaly "
     "and AI layers)."),
    ("4.2  Agentic AI Orchestration",
     "A state machine orchestrates five specialized agents — Detection, Investigation, Threat Intelligence, "
     "Risk, and Response. Each stage runs in its own database session, streams its status to the UI over "
     "WebSocket, and fails gracefully: one failed stage never crashes the pipeline. The Investigation Agent "
     "can only call an allowlist of read-mostly tools such as event search, entity history, IP reputation, "
     "RAG retrieval, MITRE mapping and attack-graph construction."),
    ("4.3  MITRE ATT&CK Mapping",
     "Every investigation maps observed behavior to MITRE ATT&CK techniques and tactics (e.g. T1110 brute "
     "force, T1078 valid accounts, T1548 privilege escalation), so analysts instantly understand the "
     "adversary's playbook and the platform can report detection coverage and gaps per tactic."),
    ("4.4  Attack Graph Reconstruction",
     "Related entities — users, devices, IPs, credentials — are assembled into a layered attack graph that "
     "shows how an adversary moved through the environment, which assets were touched, and the blast radius "
     "of the campaign."),
    ("4.5  Explainable Risk Engine",
     "A weighted scoring model turns raw evidence into a 0–100 risk score with five transparent factors: "
     "30% behavioral anomaly, 20% threat intelligence, 20% asset criticality, 15% attack progression, and "
     "15% historical evidence. Every factor ships with its supporting evidence string."),
    ("4.6  RAG Knowledge Base",
     "A retrieval-augmented generation pipeline searches vector embeddings of playbooks, policies, CVE "
     "references and MITRE material so agents ground their summaries in organizational knowledge — and "
     "degrades gracefully if the vector store is unavailable."),
    ("4.7  Human-in-the-Loop Response",
     "The Response Agent generates concrete recommendations (revoke sessions, force MFA, isolate endpoint, "
     "block IP, reset credentials) with impact ratings. None of them execute without explicit human "
     "approval, and execution is simulated by design — the platform is a defensive prototype, not an "
     "autonomous attack tool."),
    ("4.8  Tamper-Evident Evidence Ledger",
     "Forensic evidence is hashed (SHA-256), aggregated into Merkle-tree roots, and committed to "
     "proof-of-work style blocks on a permissioned ledger. A chain audit verifies integrity, and a one-click "
     "tamper test proves that altering a single byte is detected."),
]

for title, body in core:
    doc.add_heading(title, level=2)
    doc.add_paragraph(body)

# --------------------------------------------------------------------------
# 7. Principles
# --------------------------------------------------------------------------

doc.add_heading("5. Principles", level=1)

doc.add_paragraph(
    "CyberSentinel-X is governed by a small set of firm design principles. They are not features; they are "
    "the rules that every feature must obey."
)

principles = [
    ("1. Human-in-the-loop autonomy.",
     "The AI does the heavy lifting, but a human decides. High-impact actions are queued for approval and "
     "nothing destructive is ever automatic."),
    ("2. Explainability over magic.",
     "No black boxes. Every verdict, score and recommendation cites the evidence and weights behind it, and "
     "AI agents expose concise operational summaries — never hidden chain-of-thought."),
    ("3. Safety by construction.",
     "Agents are confined to allowlisted read-mostly tools: no shell execution, no arbitrary code, no "
     "unrestricted file access, no external system reach. The platform can detect and simulate, but it "
     "cannot attack."),
    ("4. Evidence over hallucination.",
     "The default LLM provider is deterministic and evidence-grounded, so the platform works with zero API "
     "keys and never invents findings. Real LLMs (OpenAI/Gemini) are optional and fall back to the "
     "deterministic engine on any failure."),
    ("5. Security by design.",
     "bcrypt password hashing, JWT access/refresh tokens, RBAC on every route, Pydantic validation, "
     "parameterized queries, rate limiting, CSP/XSS/clickjacking headers, CORS allowlists, and an audit "
     "log that never stores secrets."),
    ("6. Honest provenance.",
     "Simulated data, dataset-derived data, and model predictions are labelled and never conflated. The "
     "threat-intel panel honestly reports when no live source is configured."),
    ("7. Modularity and graceful degradation.",
     "Every external dependency — LLM, threat-intel feed, vector store — is an abstraction behind an "
     "interface, so the system keeps working (with reduced capability) when any of them is unavailable."),
    ("8. Defensive boundary.",
     "All attacks are synthetic and simulated. The platform never touches real systems, never performs "
     "destructive real-world actions, and clearly documents this boundary."),
]

for title, body in principles:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(title + " ")
    r.bold = True
    p.add_run(body)

# --------------------------------------------------------------------------
# 8. Sector Applications
# --------------------------------------------------------------------------

doc.add_heading("6. Sectors Where CyberSentinel-X Can Be Used", level=1)

doc.add_paragraph(
    "Because CyberSentinel-X is built around open standards (MITRE ATT&CK, STIX-shaped intelligence, "
    "standard event ingestion, RBAC, audit logs) and runs as an agentic SOC, it can be deployed wherever "
    "security telemetry needs fast, explainable triage:"
)

sectors = [
    ("Government & Defence",
     "National CERTs, state data centres and defence networks: continuous monitoring, campaign "
     "correlation, and tamper-evident evidence for cyber forensics and reporting."),
    ("Banking, Finance & Fintech",
     "Detecting account takeover, fraud campaigns, credential-stuffing and data exfiltration across "
     "UPI/banking infrastructure; regulator-ready audit trails."),
    ("Healthcare",
     "Protecting hospital networks, EHR/patient data and medical devices from ransomware — with "
     "compliance-friendly reporting (e.g. India's DPDP Act alignment)."),
    ("Energy & Critical Infrastructure",
     "Power grids, water systems, oil & gas and smart-city OT/IT networks, where a fast, explainable "
     "response to targeted attacks is a matter of national security."),
    ("Telecom & ISPs",
     "Monitoring network core and subscriber-facing services for DDoS, fraud and subscriber-data theft."),
    ("E-commerce & Retail",
     "Fighting bot attacks, payment fraud, credential stuffing and API abuse on customer-facing platforms."),
    ("Education & Research",
     "University and research-lab networks that hold valuable intellectual property and large user bases."),
    ("Enterprises & MSSPs",
     "Internal enterprise SOCs and managed security service providers that run the platform as a "
     "multi-tenant offering, scaling analyst coverage with AI."),
]

make_table(doc, ["Sector", "Use Case"], sectors, widths=[1.7, 4.5])
spacer(doc)

doc.add_paragraph(
    "In every sector the value proposition is the same: the AI absorbs the flood of alerts, investigates "
    "the real incidents, explains its reasoning, and escalates only what needs a human — so a small team "
    "defends a large enterprise."
)

# --------------------------------------------------------------------------
# 9. What We Need To Study
# --------------------------------------------------------------------------

doc.add_heading("7. What We Need To Study", level=1)

doc.add_paragraph(
    "Building and advancing CyberSentinel-X requires a broad, multidisciplinary study programme. The areas "
    "below map directly to the components of the platform and are the roadmap we are following:"
)

make_table(doc,
    ["Area", "Why it matters for CyberSentinel-X"],
    [
        ["Cyber threat landscape", "Attack types (ransomware, phishing, brute force, exfiltration, privilege escalation), kill chains and modern adversary TTPs."],
        ["Machine learning & anomaly detection", "Isolation Forest, unsupervised anomaly scoring, feature engineering, model drift (PSI/KL divergence), avoiding false-positive explosions."],
        ["LLMs & agentic AI", "Prompt design, tool-calling frameworks, function allowlists, hallucination control, deterministic fallbacks, safe agent orchestration."],
        ["MITRE ATT&CK", "Technique/tactic taxonomy, detection coverage analysis, mapping observables to adversary behavior."],
        ["Retrieval-Augmented Generation (RAG)", "Vector embeddings, similarity search, chunking, knowledge-base grounding over playbooks and policies."],
        ["Network & systems security", "IP reputation, indicators of compromise, event/syslog formats, WAF concepts, secure API design (OWASP Top 10)."],
        ["Real-time systems", "WebSockets, event streaming, pub/sub patterns, live dashboards, latency budgets."],
        ["Blockchain & digital forensics", "Hash chains, Merkle trees, proof-of-work, tamper evidence, evidence chain-of-custody for legal admissibility."],
        ["Security engineering", "RBAC, JWT, bcrypt, OAuth/SSO, audit logging, secrets management, secure deployment (Docker, CI/CD)."],
        ["Project management & SIH craft", "Problem-statement analysis, demo scripting, judge Q&A, honest labelling of simulated vs real data."],
    ],
    widths=[1.9, 4.3],
)

# --------------------------------------------------------------------------
# 10. Dataset section
# --------------------------------------------------------------------------

doc.add_heading("8. About the Dataset We Researched — UNSW-NB15", level=1)

doc.add_paragraph(
    "To give CyberSentinel-X a realistic, large-scale ground truth, we researched and integrated the "
    "UNSW-NB15 network-traffic dataset from the Australian Centre for Cyber Security (ACCS) at the "
    "University of New South Wales (UNSW Canberra). Released in 2015 by Nour Moustafa and Jill Slay, it is "
    "one of the most widely cited benchmarks for modern network intrusion-detection research, and it is "
    "the dataset that powers the real (non-simulated) data inside our SOC console."
)

doc.add_heading("8.1  What the dataset is", level=2)
doc.add_paragraph(
    "UNSW-NB15 is a labelled corpus of modern network traffic. The raw packets were generated with the "
    "IXIA PerfectStorm tool in the Cyber Range Lab of UNSW Canberra, producing a hybrid of real modern "
    "normal activity and synthetic contemporary attack behaviour. tcpdump captured roughly 100 GB of raw "
    "traffic (pcap files), and the Argus and Bro-IDS tools, together with twelve purpose-built algorithms, "
    "were used to derive 49 features plus a class label for every flow."
)

bullets(doc, [
    ("Scale: ", "2,540,044 records in total, spread across four CSV files (UNSW-NB15_1 to _4), with a ground-truth table and an event list."),
    ("Features: ", "49 features per flow — flow features (protocol, state, duration, packet/byte counts), basic features, content features (ct_* connection counters), time features, and additional generated features."),
    ("Attack families: ", "nine labelled attack categories plus normal traffic — Fuzzers, Analysis, Backdoors, DoS, Exploits, Generic, Reconnaissance, Shellcode and Worms."),
    ("Official split: ", "175,341 records in the training set and 82,332 in the testing set — 257,673 flows in total, which is exactly the corpus we ingested."),
    ("Why it is modern: ", "unlike the older KDD99 / NSL-KDD sets, it includes low-footprint contemporary attack behaviours (backdoors, shellcode, worms, fuzzers, exploits) that resemble today's real threat landscape."),
])

doc.add_heading("8.2  How we used it in CyberSentinel-X", level=2)
doc.add_paragraph(
    "UNSW-NB15 is not a side exhibit — it is the real data pipeline of the platform. Our ingestion service "
    "parses the official training and testing CSVs, maps every flow into the platform's event model, runs "
    "the hybrid detection engine over the corpus, and auto-correlates the results into alerts and incidents "
    "so the SOC console lights up with genuine, labelled attack data."
)

bullets(doc, [
    ("Full-corpus ingestion: ", "257,673 flows ingested (164,673 attack flows, 93,000 benign) in our reference deployment, with progress tracking and safe re-runs."),
    ("Event mapping: ", "each attack family is mapped to an event type and severity (e.g. Backdoor, Exploits, Shellcode and Worms are scored CRITICAL) and fed through the same detection, investigation and risk pipeline as live telemetry."),
    ("Feature preservation: ", "all 49 raw UNSW features are kept in the event metadata, enabling the 3D threat-space visualisation and per-family analytics on the dashboard."),
    ("IP synthesis: ", "the public CSVs ship without IP columns, so deterministic source/destination IPs are generated per row to keep the attack graph and entity analytics meaningful."),
    ("Hybrid scoring: ", "an Isolation Forest model is fit on a stratified sample of the corpus and combined with attack-family rules, so every flow receives an explainable anomaly score."),
    ("Auto-correlation: ", "attack flows are aggregated into alerts — one per family — and the most severe families are escalated into full incidents with investigations, MITRE mappings and reports."),
    ("Honest provenance: ", "everything derived from the dataset is labelled DATASET in the UI, and model outputs are labelled MODEL PREDICTION — simulated data is never mixed with dataset data."),
])

doc.add_heading("8.3  Attack families and our severity mapping", level=2)
make_table(doc,
    ["UNSW-NB15 Attack Family", "Mapped Event Type (CyberSentinel-X)", "Severity"],
    [
        ["Reconnaissance", "PORT_SCAN", "LOW"],
        ["Fuzzers", "SUSPICIOUS_NETWORK_CONNECTION", "MEDIUM"],
        ["Analysis", "SUSPICIOUS_PROCESS", "MEDIUM"],
        ["Backdoor", "MALWARE_DETECTED", "CRITICAL"],
        ["DoS", "SUSPICIOUS_NETWORK_CONNECTION", "HIGH"],
        ["Exploits", "PRIVILEGE_ESCALATION", "CRITICAL"],
        ["Generic", "SUSPICIOUS_NETWORK_CONNECTION", "MEDIUM"],
        ["Shellcode", "MALWARE_DETECTED", "CRITICAL"],
        ["Worms", "MALWARE_DETECTED", "CRITICAL"],
    ],
    widths=[2.2, 2.7, 1.3],
)
spacer(doc)

doc.add_heading("8.4  Why we chose UNSW-NB15", level=2)
bullets(doc, [
    ("Modern attack variety — ", "covers nine contemporary attack families instead of the legacy attacks in KDD99/NSL-KDD, so our detection results are relevant to today's threats."),
    ("Labelled ground truth — ", "every record is labelled normal or one of nine attacks, giving us reliable precision, recall and MTTD/MTTR numbers to report."),
    ("Realistic scale — ", "2.5 million flows / 100 GB of traffic stress the ingestion, anomaly-scoring and correlation pipeline the way production telemetry would."),
    ("Benchmark comparability — ", "it is the standard benchmark for modern intrusion-detection systems, so judges and reviewers can compare our results with published literature."),
    ("Free for academic use — ", "freely available for research; the authors request citation of the Moustafa & Slay publications."),
])

doc.add_heading("8.5  Dataset preview — rows, columns and values", level=2)
doc.add_paragraph(
    "The screenshot below is generated from the real UNSW-NB15 training file we bundled with the project "
    "(unsw_sample.csv). It shows a slice of the dataset as it sits on disk — 8 of the 175,341 rows and 16 "
    "of the 45 columns — spanning the feature groups the platform ingests: flow and basic features "
    "(id, dur, proto, service, state, packet and byte counts, rate, TTL), connection features (tcprtt, "
    "smean), a derived content feature (ct_state_ttl), and the ground-truth columns every model trains on "
    "— attack_cat (the nine attack families plus Normal) and label (0 = benign, 1 = attack)."
)

dataset_shot = SHOTS / "06-unsw-dataset.png"
if dataset_shot.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.add_run().add_picture(str(dataset_shot), width=Inches(6.2))
    caption(doc, "Figure 8.5 — UNSW-NB15 training set preview (unsw_sample.csv): 8 rows × 16 of 45 columns; "
                 "green = benign flows, red = attack flows (Backdoor, Analysis, Fuzzers, Shellcode), with "
                 "the binary label in the final column.")

doc.add_heading("8.6  Full 45-column schema", level=2)
doc.add_paragraph(
    "The complete header of the training file, showing all 45 columns the platform ingests — flow and "
    "basic features, content features (ct_* connection counters), time and additional generated features, "
    "plus the two ground-truth columns (attack_cat, label). Four sample rows are included: one benign flow "
    "and the first flows of three attack families. The wide image is best read by zooming into the page."
)

full_hdr_shot = SHOTS / "07-unsw-full-header.png"
if full_hdr_shot.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.add_run().add_picture(str(full_hdr_shot), width=Inches(6.2))
    caption(doc, "Figure 8.6 — UNSW-NB15 training set: the full 45-column header with sample rows "
                 "(1 benign + Backdoor, Analysis, Fuzzers).")

doc.add_heading("8.7  Attack-category distribution", level=2)
doc.add_paragraph(
    "The class distribution of the 175,341 training flows, computed from the bundled file. It shows the "
    "challenge our detection engine faces: the dataset is heavily imbalanced — Normal dominates at 31.9%, "
    "while the rarest families (Worms 0.07%, Shellcode 0.65%) have very few samples. This is exactly why "
    "CyberSentinel-X combines an unsupervised Isolation Forest (which does not need balanced labels) with "
    "category rules and an explainable risk engine, and why the platform tracks model drift and data "
    "quality over time."
)

dist_shot = SHOTS / "08-unsw-class-distribution.png"
if dist_shot.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.add_run().add_picture(str(dist_shot), width=Inches(5.6))
    caption(doc, "Figure 8.7 — UNSW-NB15 training set class distribution (175,341 rows): Normal 56,000 "
                 "(31.9%), Generic 40,000 (22.8%), Exploits 33,393 (19.0%), Fuzzers 18,184 (10.4%), "
                 "DoS 12,264 (7.0%), Reconnaissance 10,491 (6.0%), Analysis 2,000 (1.1%), Backdoor 1,746 "
                 "(1.0%), Shellcode 1,133 (0.7%), Worms 130 (0.1%).")

# --------------------------------------------------------------------------
# 11. Screenshots
# --------------------------------------------------------------------------

doc.add_heading("9. Project Screenshots", level=1)

doc.add_paragraph(
    "The following screenshots show the working CyberSentinel-X application as demonstrated for SIH 2026. "
    "Each screen is a live page of the running platform."
)

shots = [
    ("01-login.png", "Screenshot 1 — Secure Login: role-based access with email/password and optional "
                     "enterprise single sign-on (Google/GitHub) via OAuth."),
    ("02-dashboard.png", "Screenshot 2 — SOC Dashboard: live KPI cards (alerts, criticals, active incidents, "
                         "high risk, anomalies, approvals), AI agent status, and real-time risk and severity charts."),
    ("03-defense-center.png", "Screenshot 3 — Defense Center: layered web-application firewall with "
                              "rate limiting, payload inspection, malware-hash blocking and brute-force protection."),
    ("04-admin-users.png", "Screenshot 4 — Admin & Users: RBAC management of ADMIN, SECURITY_ANALYST and "
                           "VIEWER roles with full auditability."),
    ("05-settings.png", "Screenshot 5 — Settings: platform configuration, sign-in methods and security options."),
]

for fname, cap in shots:
    path = SHOTS / fname
    if not path.exists():
        doc.add_paragraph(f"[missing screenshot: {fname}]")
        continue
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.add_run().add_picture(str(path), width=Inches(6.0))
    caption(doc, cap)

# --------------------------------------------------------------------------
# 11. Conclusion & Future Scope
# --------------------------------------------------------------------------

doc.add_heading("10. Conclusion", level=1)

doc.add_paragraph(
    "CyberSentinel-X demonstrates a practical, trustworthy vision of AI-driven cybersecurity: an agentic "
    "platform that detects, investigates, explains, responds and documents — with the discipline of a "
    "well-run SOC built in. Its hybrid detection avoids the classic AI pitfall of hallucinated alarms; its "
    "explainable risk engine and evidence ledger make every decision auditable; and its human-in-the-loop "
    "response design keeps control where it belongs. For Smart India Hackathon 2026, it is a complete, "
    "runnable, test-backed implementation of that vision — not a slideware concept."
)

doc.add_heading("11. Future Scope", level=1)
bullets(doc, [
    "Live threat-intelligence integration: STIX/TAXII feeds and vendor APIs replacing the local feed.",
    "Scalable vector backends (pgvector / FAISS) and streaming analytics for very high event rates.",
    "Multi-tenancy and SSO federation for MSSP and government deployments.",
    "XGBoost / deep-learning classifiers for attack-type prediction on top of the anomaly engine.",
    "Real EDR/SIEM ingestion connectors (syslog, CEF, OCSF) so the platform consumes production telemetry.",
    "SOAR-style playbook automation with external security integrations, still gated by human approval.",
    "Privacy-preserving federated learning across participating organizations (e.g. across banks or states).",
    "Deployable zero-trust agent sensors to extend detection to endpoints and cloud workloads.",
])

doc.save(OUT)
print(f"Saved: {OUT}")
